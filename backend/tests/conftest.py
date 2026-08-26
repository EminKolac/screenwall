"""Golden fixtures: synthesize TR / EN / mixed DOCX, XLSX, and (English) PDF in-memory.

PDF fixtures use English text only: synthetic generation with base-14 fonts cannot embed Turkish
glyphs (ş/ğ/ı). Real Turkish PDFs embed fonts and extract fine; TR/mixed detection is therefore
exercised via DOCX/XLSX (full Unicode) where synthetic generation is faithful.
"""
from __future__ import annotations

import io

import openpyxl
import pytest
from docx import Document as Docx


@pytest.fixture(autouse=True)
def _isolated_storage(tmp_path, monkeypatch):
    """Point every test at a temp STORAGE_ROOT so the storage-backed repo never touches ./data.

    Also pins the auditor to the deterministic heuristic ("mlx" — no LLM). Without this, tests pass
    or fail depending on whether an unrelated local Ollama server happens to be running: the Qwen
    auditor is far stricter than the heuristic (approved vs. 3 iterations -> needs_human_review on
    the exact same input), so results must not depend on the environment.
    """
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path))
    monkeypatch.setenv("AUDITOR_PROVIDER", "mlx")
    from app.anonymization.privacy_filter import get_privacy_filter
    from app.config import get_settings
    from app.services.deps import get_repository

    get_settings.cache_clear()
    get_repository.cache_clear()
    get_privacy_filter.cache_clear()  # settings-dependent; must not leak across tests
    yield
    get_settings.cache_clear()
    get_repository.cache_clear()
    get_privacy_filter.cache_clear()


TR_PARA = "Bu sözleşme Türkiye Cumhuriyeti kanunlarına tabidir ve taraflar arasında imzalanmıştır."
TR_PARA2 = "Yatırımcı, fona katılım payı taahhüt etmiş olup ödemeleri zamanında yapacaktır."
EN_PARA = "This agreement is governed by the laws of England and Wales and binds both parties."
EN_PARA2 = "The investor shall fund all capital calls within ten business days of each notice."


def make_docx(paragraphs: list[str], heading: str | None = None,
              table: list[list[str]] | None = None) -> bytes:
    d = Docx()
    if heading:
        d.add_heading(heading, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def make_xlsx(rows: list[list[str]], sheet: str = "Sheet1") -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_pdf(lines: list[str]) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontsize=12)
        y += 22
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def docx_tr() -> bytes:
    return make_docx([TR_PARA, TR_PARA2], heading="Sözleşme",
                     table=[["Ad", "Tutar"], ["Ahmet Yılmaz", "1000"]])


@pytest.fixture
def docx_en() -> bytes:
    return make_docx([EN_PARA, EN_PARA2], heading="Agreement",
                     table=[["Name", "Amount"], ["John Smith", "1000"]])


@pytest.fixture
def docx_mixed() -> bytes:
    return make_docx([TR_PARA, EN_PARA, TR_PARA2, EN_PARA2], heading="Side Letter")


@pytest.fixture
def xlsx_tr() -> bytes:
    return make_xlsx([["Ad", "Tutar"], ["Ahmet Yılmaz", "1000"], ["Mehmet Demir", "2000"]])


@pytest.fixture
def pdf_en() -> bytes:
    return make_pdf([
        "Confidential Due Diligence Report",
        "This document was prepared by the finance team for the investor.",
        "The total committed amount is 5000 USD across all parties.",
    ])
