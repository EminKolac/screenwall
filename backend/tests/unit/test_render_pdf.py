"""v1 export: generated-PDF renderer + full-coverage DOCX/XLSX extraction + OCR fail-closed.

The download is built ONLY from the audited anonymized content (layer 3), so these tests assert
that (a) rendering preserves placeholders + Turkish glyphs + structure, (b) DOCX headers/footers/
comments and XLSX cell comments are now extracted (were leak channels), and (c) a page that needs
OCR but has no Tesseract fails closed to human review instead of emitting partial text.
"""
from __future__ import annotations

import io
import shutil

import fitz
import openpyxl
import pytest
from docx import Document as Docx
from openpyxl.comments import Comment

from app.export.render_pdf import _FONTS_DIR, render_content_pdf
from app.extraction.base import Block, BlockType, ExtractedContent, TableCell
from app.extraction.docx import DocxExtractor
from app.extraction.xlsx import XlsxExtractor
from app.models.document import FileKind


def _pdf_text(pdf: bytes) -> str:
    d = fitz.open(stream=pdf, filetype="pdf")
    try:
        return "\n".join(p.get_text() for p in d)
    finally:
        d.close()


# ------------------------------- renderer ----------------------------------

def test_render_preserves_tokens_turkish_and_structure():
    anon = ExtractedContent(kind=FileKind.docx, blocks=[
        Block(block_id="h", type=BlockType.heading, text="Yatırımcı Özeti çğıİöşü"),
        Block(block_id="p", type=BlockType.paragraph, text="Sorumlu <PERSON_1>, TCKN <TCKN_1>."),
        Block(block_id="hd", type=BlockType.paragraph, text="<ORG_1> gizli", location="header"),
        Block(block_id="ft", type=BlockType.paragraph, text="footer <PERSON_2>", location="footer"),
        Block(block_id="cm", type=BlockType.paragraph, text="yorum <PHONE_1>", location="comment"),
        Block(block_id="t", type=BlockType.table, sheet="Sheet1", cells=[
            TableCell(row=0, col=0, text="Ad"), TableCell(row=0, col=1, text="Tutar"),
            TableCell(row=1, col=0, text="<PERSON_3>"), TableCell(row=1, col=1, text="1.250.000")]),
    ])
    pdf = render_content_pdf(anon)
    assert pdf[:4] == b"%PDF"
    txt = _pdf_text(pdf)
    for token in ("<PERSON_1>", "<TCKN_1>", "<ORG_1>", "<PHONE_1>", "<PERSON_3>"):
        assert token in txt
    assert "Yatırımcı" in txt and "çğıİöşü" in txt  # Turkish glyphs render (embedded font)
    for label in ("[Header]", "[Footer]", "[Comment]"):
        assert label in txt


def test_anonymize_filename_masks_pii():
    from app.export.filename import anonymize_filename
    from app.models.document import Language

    out = anonymize_filename("Rapor ahmet@example.com Q1.pdf", Language.tr)
    assert "ahmet@example.com" not in out       # email (deterministic recognizer) masked
    assert "EMAIL" in out                         # -> EMAIL_n token (brackets stripped)
    for bad in (" ", "/", "<", ">", "@"):
        assert bad not in out                     # filesystem-safe
    # deny-listed brand names (NER can't catch them) must be masked in the filename too
    out2 = anonymize_filename("e2vc Fund III Deck.pdf", Language.tr, deny_terms=["e2vc"])
    assert "e2vc" not in out2


def test_render_empty_content_is_valid_pdf():
    pdf = render_content_pdf(ExtractedContent(kind=FileKind.pdf, blocks=[]))
    assert pdf[:4] == b"%PDF"
    assert "no extractable content" in _pdf_text(pdf)


def test_render_output_is_subset_and_small():
    anon = ExtractedContent(kind=FileKind.pdf,
                            blocks=[Block(block_id="p", text="Türkçe çğıİöşü <PERSON_1>")])
    assert len(render_content_pdf(anon)) < 300_000  # font subsetted, not the full ~1.5 MB DejaVu


# --------------------------- DOCX full coverage ----------------------------

def test_docx_extracts_header_footer_comment():
    d = Docx()
    p = d.add_paragraph("Gövde metni: Ahmet Yılmaz")
    d.sections[0].header.paragraphs[0].text = "HDR Vakıf Holding"
    d.sections[0].footer.paragraphs[0].text = "FTR gizli 0532 000 1122"
    d.add_comment(runs=p.runs, text="CMT john@acme.com", author="rev", initials="r")
    buf = io.BytesIO()
    d.save(buf)

    content = DocxExtractor().extract(buf.getvalue(), "x.docx")
    by_loc = {b.location: b.text for b in content.blocks if b.type == BlockType.paragraph}
    assert any(b.location == "header" and "HDR Vakıf" in b.text for b in content.blocks)
    assert any(b.location == "footer" and "FTR gizli" in b.text for b in content.blocks)
    assert any(b.location == "comment" and "john@acme.com" in b.text for b in content.blocks)
    assert "Gövde metni: Ahmet Yılmaz" in " ".join(by_loc.values()) or any(
        "Ahmet" in b.text for b in content.blocks)


# --------------------------- XLSX full coverage ----------------------------

def test_xlsx_extracts_cell_comment():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "S1"
    ws["A1"] = "Ad"
    ws["A2"] = "Ahmet Yılmaz"
    ws["A2"].comment = Comment("Not: 0532 111 2233", "rev")
    buf = io.BytesIO()
    wb.save(buf)

    content = XlsxExtractor().extract(buf.getvalue(), "x.xlsx")
    assert any(b.type == BlockType.table and b.sheet == "S1" for b in content.blocks)
    # Sheet name is also emitted as a heading block so it gets anonymized + audited (not raw).
    assert any(b.type == BlockType.heading and b.text == "S1" for b in content.blocks)
    assert any(b.location == "comment" and "0532 111 2233" in b.text for b in content.blocks)


# ------------------------------ OCR fail-closed ----------------------------

def _image_only_pdf(text: str) -> bytes:
    """A 1-page PDF whose text is rendered as an IMAGE (no text layer) — needs OCR to read."""
    src = fitz.open()
    page = src.new_page()
    page.insert_text((36, 80), text, fontsize=18,
                     fontfile=str(_FONTS_DIR / "DejaVuSans.ttf"), fontname="djv")
    pix = page.get_pixmap(dpi=150)
    out = fitz.open()
    p2 = out.new_page(width=page.rect.width, height=page.rect.height)
    p2.insert_image(p2.rect, pixmap=pix)
    data = out.tobytes()
    src.close()
    out.close()
    return data


@pytest.mark.skipif(not shutil.which("tesseract"), reason="tesseract not installed")
def test_image_pdf_ocr_recovers_text():
    """Image-only PDF (no text layer) → OCR at extraction makes the baked text visible again.

    This is the core of the Türk Telekom fix: without OCR the text is invisible to anonymization.
    Assertion is loose (OCR is approximate) — it only proves OCR ran and recovered the content.
    """
    data = _image_only_pdf("Turk Telekomunikasyon A.S. Tel 0532 111 2233")
    from app.extraction.pdf import PdfExtractor

    content = PdfExtractor().extract(data, "scan.pdf")
    txt = content.plain_text
    assert "Telekom" in txt and "0532" in txt  # OCR recovered the image text


def test_ocr_unavailable_routes_to_human_review(monkeypatch):
    import app.extraction.pdf as pdf_mod
    from app.config import Settings
    from app.models.document import DocumentStatus
    from app.services import ingest as ingest_mod

    monkeypatch.setattr(pdf_mod, "_tesseract_available", lambda: False)
    data = _image_only_pdf("Türk Telekom A.Ş. iletisim@example.com")
    doc, content = ingest_mod.ingest(data, "scan.pdf", Settings())
    assert doc.status == DocumentStatus.NEEDS_HUMAN_REVIEW
    from app.extraction.base import OCR_UNAVAILABLE
    assert any(w.startswith(OCR_UNAVAILABLE) for w in content.warnings)
