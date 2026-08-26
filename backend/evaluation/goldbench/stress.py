"""GoldBench stres korpusu — PII recall'unu DEĞİL, güvenlik kapılarını ölçer.

Buradaki 72 belge (24 pdf / 24 docx / 24 xlsx) "sistemin göremediği ya da göremeyeceği" yüzeylere
sentetik PII yerleştirir. Her belgenin beklentisi İKİLİ:

  safe_output  — belge güvenle işlendi; yerleştirilen değerlerin hiçbiri anonim çıktıda veya
                 export edilen PDF'te görünmüyor. (Belgenin NEEDS_HUMAN_REVIEW'a düşmesi de
                 kabul edilir — temkinli ama güvenli; `run_stress` bunu `over_conservative`
                 olarak ayrıca işaretler.)
  fail_closed  — sistem riski fark edip belgeyi reddetmeli VEYA insan incelemesine yönlendirmeli.

Tek kritik ihlal: desteklenmeyen/riskli bir yüzey varken APPROVED + sızıntı. Bunu `run_stress.py`
`critical_false_approval` olarak sayar ve release gate 0 bekler.

MİMARİ NOT — "test edildi ve geçti" ile "mimari gereği mümkün değil" aynı şey DEĞİLDİR.
Bu sistemin export'u orijinal dosyayı yeniden serileştirmez; layer-3 (anonim `ExtractedContent`)
üzerinden sıfırdan yeni bir PDF üretir (`app/export/render_pdf.py`). Dolayısıyla `metadata` ve
`external_link` gibi kanallar export'a hiçbir koşulda taşınamaz — bunlar `architecturally_safe=True`
ile işaretlenir ve raporda AYRI gösterilir. Bir güvenceyi kanıtlayan şey testin geçmesi değil,
kanalın var olmamasıdır; ikisini aynı sütunda toplamak sahte güven üretir.

DETERMİNİZM: sabit seed + sabit belge tarihleri + normalize edilmiş konteynerler → aynı byte'lar,
aynı sha256. PDF'te trailer `/ID` her kayıtta yeniden üretildiği için uzunluk koruyarak sıfırlanır
(`_freeze_pdf`); docx/xlsx'te `created/modified` sabitlenir. `corpus_hashes()` bunu doğrular.

GÜVENLİK: buradaki değerler `evaluation/bist30/canary.py` kataloğundaki SENTETİK değerlerdir
(gerçek kişi/kurum yok). Yine de rapora/JSONL'e ham değer yazılmaz — sadece sha256[:16].
"""
from __future__ import annotations

import hashlib
import io
import re
import zipfile
from dataclasses import asdict, dataclass, field
from datetime import datetime

# Sabit seed: üretim sırası, değer rotasyonu ve varyant seçimi bundan türetilir.
SEED = 20260812

# Tüm konteynerlerde kullanılan sabit belge tarihi (determinizm için).
FIXED_DT = datetime(2024, 1, 1, 0, 0, 0)
_PDF_DATE = "D:20240101000000Z"

FORMATS = ("pdf", "docx", "xlsx")
PER_FORMAT = 24

SAFE_OUTPUT = "safe_output"
FAIL_CLOSED = "fail_closed"

# Bir senaryonun hangi formatlarda ANLAMLI olduğu. Anlamsız kombinasyon üretilmez
# (örn. XLSX'te "pdf_form_annotation" diye bir yüzey yoktur).
SCENARIO_FORMATS: dict[str, tuple[str, ...]] = {
    "split_run_pii": ("pdf", "docx", "xlsx"),
    "header_footer": ("pdf", "docx", "xlsx"),
    "comment": ("docx", "xlsx"),
    "hidden_sheet": ("xlsx",),
    "sheet_name": ("xlsx",),
    "metadata": ("pdf", "docx", "xlsx"),
    "external_link": ("pdf", "docx", "xlsx"),
    "scanned_image": ("pdf",),
    "pdf_form_annotation": ("pdf",),
    "corrupt_ooxml": ("docx", "xlsx"),
    "zip_bomb": ("docx", "xlsx"),
    "format_variants": ("pdf", "docx", "xlsx"),
}

# Export orijinali yeniden serileştirmediği için bu kanallar mimari olarak çıktıya ULAŞAMAZ.
ARCHITECTURALLY_SAFE = frozenset({"metadata", "external_link"})

# Konteynerin kendisi bozuk/tehlikeli → `validate_upload` fail-closed reddetmeli.
_FAIL_CLOSED_SCENARIOS = frozenset({"corrupt_ooxml", "zip_bomb"})

SCENARIO_NOTES: dict[str, str] = {
    "split_run_pii": "Değer satır sonuna / ayrı run'lara / bitişik hücrelere bölündü.",
    "header_footer": "PII gövdede değil header/footer bandında (xlsx: sayfa üstbilgisi).",
    "comment": "PII belge yorumunda — gövde metnine hiç girmiyor.",
    "hidden_sheet": "PII gizli (hidden) sayfada.",
    "sheet_name": "PII sayfa adının kendisinde.",
    "metadata": "PII belge metadata'sında (author/subject/creator).",
    "external_link": "PII dış bağlantı URL'inde (ilişki/hyperlink hedefi).",
    "scanned_image": "PII sayfa görseline gömülü — okunması OCR gerektirir.",
    "pdf_form_annotation": "PII PDF form alanı (widget) ve freetext annotation içinde.",
    "corrupt_ooxml": "Geçerli zip magic, bozuk OOXML paketi → reddedilmeli.",
    "zip_bomb": "Aşırı sıkıştırma oranı → zip-bomb koruması devreye girmeli.",
    "format_variants": "Aynı değerin boşluklu/boşluksuz/büyük harfli yazımları bir arada.",
}


@dataclass
class StressCase:
    """Tek bir stres belgesinin tanımı ve beklentisi.

    `generated=False` olan vaka SESSİZCE düşürülmez: gerekçesiyle korpusta kalır ve raporda
    ayrı sayılır — "üretilemedi" ile "üretildi ve geçti" karıştırılamaz.
    """

    case_id: str
    fmt: str
    scenario: str
    expected: str                                  # safe_output | fail_closed
    planted_values: list[str] = field(default_factory=list)  # ham SENTETİK PII — rapora girmez
    context: str = "Bilgi"                         # değerin yanındaki PII OLMAYAN bağlam kelimesi
    note: str = ""
    architecturally_safe: bool = False
    generated: bool = True
    skip_reason: str = ""

    def safe_dict(self) -> dict:
        """Rapor-güvenli görünüm — ham değer YOK, sadece sha256[:16] listesi."""
        d = asdict(self)
        d.pop("planted_values", None)
        d["planted_vhashes"] = [
            hashlib.sha256(v.encode("utf-8")).hexdigest()[:16] for v in self.planted_values
        ]
        return d


# ----------------------------------------------------------------------------------------------
# Değer havuzu — bist30 sentetik kataloğundan, sabit sırayla.
# ----------------------------------------------------------------------------------------------

def _value_pool() -> list[tuple[str, str]]:
    """(bağlam kelimesi, değer) çiftleri. Katalog sabit olduğu için sıra da sabit.

    Bağlam kelimesi ("IBAN", "cep", "e-posta"…) kataloğun kendi alanıdır: düşük skorlu TR
    recognizer'ları bağlam olmadan tetiklenmez, o yüzden değer yalın bırakılmaz."""
    from evaluation.bist30.canary import canary_catalog

    return [(c.context or c.expected_family, c.value) for c in canary_catalog()]


def _pick(index: int) -> tuple[str, str]:
    """Sabit seed'li deterministik rotasyon — havuz boyutuyla aralarında asal adım kullanır."""
    pool = _value_pool()
    return pool[(index * 7 + SEED) % len(pool)]


# ----------------------------------------------------------------------------------------------
# Konteyner normalizasyonu (determinizm)
# ----------------------------------------------------------------------------------------------

def _pdf_array_end(data: bytes, start: int) -> int:
    """`start` konumundaki `[` ile eşleşen `]`in indeksi (yoksa -1).

    Basit bir `[^\\]]*` regex'i YETMEZ: `/ID`nin elemanları hex string (`<...>`) ya da literal
    string (`(...)`) olabilir ve literal string kaçışlı parantez veya `]` içerebilir. Bu tarayıcı
    her iki string biçimini de atlayarak dizinin gerçek sonunu bulur.
    """
    i, n = start + 1, len(data)
    while i < n:
        ch = data[i:i + 1]
        if ch == b"]":
            return i
        if ch == b"(":
            depth, i = 1, i + 1
            while i < n and depth:
                c = data[i:i + 1]
                if c == b"\\":
                    i += 2
                    continue
                depth += 1 if c == b"(" else (-1 if c == b")" else 0)
                i += 1
            continue
        if ch == b"<":
            end = data.find(b">", i)
            if end < 0:
                return -1
            i = end + 1
            continue
        i += 1
    return -1


def _freeze_pdf(data: bytes) -> bytes:
    """Trailer `/ID`nin ikinci elemanı her kayıtta yeniden üretilir → sabit bir değerle değiştir.

    MuPDF bu elemanı bazen hex (`<A1B2…>`), bazen literal string (`(\\013\\257…)`) olarak yazar —
    hangisinin seçileceği rastgele byte'lara bağlıdır, bu yüzden sadece hex biçimini normalize
    eden bir yaklaşım ARADA BİR kaçırır ve determinizm testini flaky yapar.

    `/ID` dosyanın sonundaki trailer/xref sözlüğündedir; sonrasında offset referansı kalmadığı
    için uzunluğun değişmesi güvenlidir (`startxref` kendinden ÖNCEKİ xref'i gösterir).
    """
    i = data.rfind(b"/ID")
    if i < 0:
        return data
    j = data.find(b"[", i)
    if j < 0 or j - i > 8:  # araya başka bir anahtar girmişse dokunma
        return data
    k = _pdf_array_end(data, j)
    if k < 0:
        return data
    return data[:i] + b"/ID[<0000000000000000><0000000000000000>]" + data[k + 1:]


_ISO_DT_RE = re.compile(rb"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}Z")
_FIXED_ISO = b"2024-01-01T00:00:00Z"


def _is_timestamped_part(name: str) -> bool:
    """Kayıt anının saatini gömen OOXML parçaları (yalnızca bunlar normalize edilir — gövde
    XML'ine dokunulmaz, aksi halde belge içeriği sessizce değişebilir)."""
    return name.startswith("docProps/") or name.startswith("word/comments")


def _freeze_zip(data: bytes) -> bytes:
    """OOXML konteynerini yeniden yaz: giriş sırası ve içerik aynı, zaman damgaları sabit.

    İki ayrı zaman kaynağı vardır ve ikisi de sabitlenmelidir:
      1. Zip girişi `date_time` — python-docx/openpyxl `writestr`'ı varsayılanla çağırır, yani
         her kayıt o anki saati gömer.
      2. Parça İÇİNDEKİ ISO damgaları — openpyxl `docProps/core.xml`'e `dcterms:modified`,
         python-docx `word/comments.xml`'e `w:date` yazar; ikisi de kayıt anında üretilir.

    Bu normalizasyon olmadan "aynı seed → aynı sha256" garantisi saniye sınırında bozulur.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as src, \
            zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as out:
        for info in src.infolist():
            body = src.read(info.filename)
            if _is_timestamped_part(info.filename):
                body = _ISO_DT_RE.sub(_FIXED_ISO, body)
            frozen = zipfile.ZipInfo(info.filename, date_time=_ZIP_DATE)
            frozen.compress_type = info.compress_type
            frozen.external_attr = info.external_attr
            frozen.create_system = info.create_system
            out.writestr(frozen, body)
    return buf.getvalue()


def _freeze_pdf_annots(doc) -> None:
    """Annotation ve form widget'ları `/M` (değiştirilme tarihi) alanını kayıt anında yazar.
    Sabitlenmezse aynı vaka her üretimde farklı byte'lar verir."""
    for page in doc:
        xrefs = [a.xref for a in page.annots()] + [w.xref for w in page.widgets()]
        for xref in xrefs:
            try:
                doc.xref_set_key(xref, "M", f"({_PDF_DATE})")
            except Exception:  # noqa: BLE001 — damga sabitlenemezse üretim yine de sürer
                pass


def _pdf_bytes(doc, extra_meta: dict[str, str] | None = None) -> bytes:
    _freeze_pdf_annots(doc)
    meta ={"title": "GoldBench Stress", "author": "", "subject": "", "keywords": "",
            "creator": "goldbench", "producer": "goldbench",
            "creationDate": _PDF_DATE, "modDate": _PDF_DATE}
    meta.update(extra_meta or {})
    doc.set_metadata(meta)
    out = doc.tobytes(garbage=4, deflate=True)
    doc.close()
    return _freeze_pdf(out)


def _font_file() -> str:
    from app.export.render_pdf import _FONTS_DIR

    return str(_FONTS_DIR / "DejaVuSans.ttf")


def _new_pdf():
    import fitz

    doc = fitz.open()
    doc.new_page()
    return doc


def _write(page, y: float, text: str, size: float = 10) -> None:
    page.insert_text((36, y), text, fontsize=size, fontfile=_font_file(), fontname="djv")


def _new_docx():
    import docx as _docx

    d = _docx.Document()
    d.core_properties.created = FIXED_DT
    d.core_properties.modified = FIXED_DT
    d.core_properties.author = "goldbench"
    d.core_properties.last_modified_by = "goldbench"
    return d


def _docx_bytes(d) -> bytes:
    buf = io.BytesIO()
    d.save(buf)
    return _freeze_zip(buf.getvalue())


def _new_xlsx():
    import openpyxl

    wb = openpyxl.Workbook()
    wb.properties.created = FIXED_DT
    wb.properties.modified = FIXED_DT
    wb.properties.creator = "goldbench"
    wb.properties.lastModifiedBy = "goldbench"
    return wb


def _xlsx_bytes(wb) -> bytes:
    buf = io.BytesIO()
    wb.save(buf)
    return _freeze_zip(buf.getvalue())


# ----------------------------------------------------------------------------------------------
# Bozuk / tehlikeli konteynerler (fail_closed beklenen)
# ----------------------------------------------------------------------------------------------

_OOXML_MAIN = {"docx": "word/document.xml", "xlsx": "xl/workbook.xml"}
_ZIP_DATE = (1980, 1, 1, 0, 0, 0)


def _corrupt_ooxml(fmt: str, salt: int) -> bytes:
    """Geçerli zip magic, geçersiz paket gövdesi → `validate_upload` BadZipFile → reddedilmeli."""
    filler = bytes((i * 31 + salt) % 256 for i in range(512))
    return b"PK\x03\x04" + filler


def _zip_bomb(fmt: str, payload_mb: int = 32) -> bytes:
    """Zorunlu OOXML parçalarını İÇEREN ama sıkıştırma oranı korumayı aşan arşiv.

    Amaç, paket-parça kontrolünü geçip `_validate_ooxml`'in zip-bomb eşiğine (ratio > 200)
    ulaşmak — yani reddin sebebinin "eksik parça" değil gerçekten bomba koruması olması.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, body in (
            ("[Content_Types].xml", b'<?xml version="1.0"?><Types/>'),
            (_OOXML_MAIN[fmt], b'<?xml version="1.0"?><root/>'),
        ):
            info = zipfile.ZipInfo(name, date_time=_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            zf.writestr(info, body)
        info = zipfile.ZipInfo("bomb.bin", date_time=_ZIP_DATE)
        info.compress_type = zipfile.ZIP_DEFLATED
        zf.writestr(info, b"\x00" * (payload_mb * 1024 * 1024))
    return buf.getvalue()


# ----------------------------------------------------------------------------------------------
# Senaryo üreticileri — PDF
# ----------------------------------------------------------------------------------------------

def _split_fragments(value: str) -> list[str]:
    """Değeri iki parçaya böl. Kısa parçalar sızıntı aramasında yanlış pozitif üretmesin diye
    çağıran taraf 6 karakterden kısa olanları eler."""
    mid = max(1, len(value) // 2)
    return [value[:mid].strip(), value[mid:].strip()]


def _leak_needles(value: str, fragments: list[str] | None = None) -> list[str]:
    out = [value]
    for f in fragments or []:
        if len(f) >= 6 and f not in out:
            out.append(f)
    return out


def _pdf_case(scenario: str, value: str, label: str):
    import fitz

    doc = _new_pdf()
    page = doc[0]
    extra: dict[str, str] = {}
    if scenario != "scanned_image":
        # `scanned_image` KASITLI olarak metin katmanı taşımaz: gerçek bir tarama gibi, tüm
        # içerik piksel. Buraya bir satır metin eklemek senaryoyu sulandırır (belge "okunabilir"
        # sayılır ve OCR kapısı hiç yoklanmaz).
        _write(page, 60, "GoldBench stres belgesi — sentetik içerik.")

    if scenario == "split_run_pii":
        a, b = _split_fragments(value)
        _write(page, 90, f"{label} bilgisi: {a}")
        _write(page, 104, b)
    elif scenario == "header_footer":
        _write(page, 24, f"{label}: {value}", size=8)                      # sayfa üstü bandı
        _write(page, page.rect.height - 18, f"{label}: {value}", size=8)   # sayfa altı bandı
    elif scenario == "metadata":
        extra = {"author": value, "subject": f"{label}: {value}", "keywords": value}
    elif scenario == "external_link":
        _write(page, 90, "Ayrıntılar için bağlantıya bakınız.")
        page.insert_link({"kind": fitz.LINK_URI, "from": fitz.Rect(36, 82, 400, 96),
                          "uri": f"https://portal.ornek.test/r?{label}={value}"})
    elif scenario == "scanned_image":
        tmp = _new_pdf()
        _write(tmp[0], 40, f"{label}: {value}", size=14)
        png = tmp[0].get_pixmap(dpi=110).tobytes("png")
        tmp.close()
        page.insert_image(fitz.Rect(36, 80, 520, 380), stream=png)
    elif scenario == "pdf_form_annotation":
        annot = page.add_freetext_annot(fitz.Rect(36, 90, 520, 120), f"{label}: {value}")
        annot.update()
        widget = fitz.Widget()
        widget.rect = fitz.Rect(36, 140, 400, 165)
        widget.field_name = "musteri_bilgi"
        widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        widget.field_value = value
        page.add_widget(widget)
    elif scenario == "format_variants":
        for i, variant in enumerate(_format_variants(value)):
            _write(page, 90 + i * 14, f"{label}: {variant}")
    else:  # pragma: no cover — SCENARIO_FORMATS ile eşleşmeyen kombinasyon
        doc.close()
        raise ValueError(f"pdf için desteklenmeyen senaryo: {scenario}")
    return _pdf_bytes(doc, extra)


# ----------------------------------------------------------------------------------------------
# Senaryo üreticileri — DOCX
# ----------------------------------------------------------------------------------------------

def _docx_case(scenario: str, value: str, label: str) -> bytes:
    d = _new_docx()
    d.add_paragraph("GoldBench stres belgesi — sentetik içerik.")

    if scenario == "split_run_pii":
        # Aynı paragrafta ayrı run'lar: çıkarım run'ları birleştirir, dedektör ise parçalı
        # gördüğü metinde değeri kaçırabilir. Klasik OOXML split-run tuzağı.
        para = d.add_paragraph()
        para.add_run(f"{label}: ")
        for chunk in _split_fragments(value):
            para.add_run(chunk)
            para.add_run(" ")
    elif scenario == "header_footer":
        section = d.sections[0]
        for target in (section.header, section.footer):
            para = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
            para.text = f"{label}: {value}"
    elif scenario == "comment":
        para = d.add_paragraph()
        run = para.add_run("Bu satır hakkında not düşüldü.")
        d.add_comment([run], text=f"{label}: {value}", author="denetci", initials="dn")
    elif scenario == "metadata":
        d.core_properties.author = value
        d.core_properties.subject = f"{label}: {value}"
        d.core_properties.keywords = value
        d.core_properties.comments = value
    elif scenario == "external_link":
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        d.add_paragraph("Ayrıntılar için bağlantıya bakınız.")
        d.part.relate_to(f"https://portal.ornek.test/r?{label}={value}",
                         RT.HYPERLINK, is_external=True)
    elif scenario == "format_variants":
        for variant in _format_variants(value):
            d.add_paragraph(f"{label}: {variant}")
    else:  # pragma: no cover
        raise ValueError(f"docx için desteklenmeyen senaryo: {scenario}")
    return _docx_bytes(d)


# ----------------------------------------------------------------------------------------------
# Senaryo üreticileri — XLSX
# ----------------------------------------------------------------------------------------------

def _xlsx_case(scenario: str, value: str, label: str) -> bytes:
    from openpyxl.comments import Comment

    wb = _new_xlsx()
    ws = wb.active
    ws.title = "Veri"
    ws["A1"] = "GoldBench stres belgesi — sentetik içerik."

    if scenario == "split_run_pii":
        # Bitişik hücreler: her hücre tek başına anlamsız, birleşince PII.
        for i, chunk in enumerate(_split_fragments(value)):
            ws.cell(row=3, column=1 + i, value=chunk)
        ws.cell(row=2, column=1, value=f"{label} (bölünmüş)")
    elif scenario == "header_footer":
        ws.oddHeader.center.text = f"{label}: {value}"
        ws.oddFooter.center.text = f"{label}: {value}"
    elif scenario == "comment":
        cell = ws.cell(row=3, column=1, value="Not eklendi")
        cell.comment = Comment(f"{label}: {value}", "denetci")
    elif scenario == "hidden_sheet":
        hidden = wb.create_sheet("Ek")
        hidden.cell(row=1, column=1, value=f"{label}: {value}")
        hidden.sheet_state = "hidden"
    elif scenario == "sheet_name":
        wb.create_sheet(_sheet_safe(f"{label} {value}"))
    elif scenario == "metadata":
        wb.properties.creator = value
        wb.properties.subject = f"{label}: {value}"
        wb.properties.keywords = value
        wb.properties.description = value
    elif scenario == "external_link":
        cell = ws.cell(row=3, column=1, value="Bağlantı")
        cell.hyperlink = f"https://portal.ornek.test/r?{label}={value}"
    elif scenario == "format_variants":
        for i, variant in enumerate(_format_variants(value)):
            ws.cell(row=3 + i, column=1, value=f"{label}: {variant}")
    else:  # pragma: no cover
        raise ValueError(f"xlsx için desteklenmeyen senaryo: {scenario}")
    return _xlsx_bytes(wb)


_SHEET_FORBIDDEN = re.compile(r"[\\/*?:\[\]]")


def _sheet_safe(name: str) -> str:
    """Excel sayfa adı kısıtı: 31 karakter, bazı karakterler yasak."""
    return _SHEET_FORBIDDEN.sub(" ", name).strip()[:31] or "Sayfa"


def _format_variants(value: str) -> list[str]:
    """Aynı değerin farklı yazımları — normalizasyon körlüğünü ölçer."""
    compact = value.replace(" ", "")
    variants = [value, compact, value.upper(), compact.lower()]
    out: list[str] = []
    for v in variants:
        if v and v not in out:
            out.append(v)
    return out


# ----------------------------------------------------------------------------------------------
# Korpus planı
# ----------------------------------------------------------------------------------------------

def _scenarios_for(fmt: str) -> list[str]:
    return [s for s in SCENARIO_FORMATS if fmt in SCENARIO_FORMATS[s]]


def plan() -> list[StressCase]:
    """Format başına tam `PER_FORMAT` vaka. Uygulanabilir senaryolar sabit sırada döngülenir,
    her tur farklı bir değer kullanır → deterministik ve dengeli dağılım."""
    cases: list[StressCase] = []
    for fmt in FORMATS:
        scenarios = _scenarios_for(fmt)
        for i in range(PER_FORMAT):
            scenario = scenarios[i % len(scenarios)]
            context, value = _pick(len(cases))
            fail_closed = scenario in _FAIL_CLOSED_SCENARIOS
            planted: list[str] = []
            if not fail_closed:
                fragments = _split_fragments(value) if scenario == "split_run_pii" else []
                if scenario == "format_variants":
                    planted = _format_variants(value)
                else:
                    planted = _leak_needles(value, fragments)
            cases.append(StressCase(
                case_id=f"{fmt}-{scenario}-{i:02d}",
                fmt=fmt,
                scenario=scenario,
                expected=FAIL_CLOSED if fail_closed else SAFE_OUTPUT,
                planted_values=planted,
                context=context,
                note=SCENARIO_NOTES[scenario],
                architecturally_safe=scenario in ARCHITECTURALLY_SAFE,
            ))
    return cases


_BUILDERS = {"pdf": _pdf_case, "docx": _docx_case, "xlsx": _xlsx_case}


def build_one(case: StressCase) -> bytes:
    """Tek bir vakanın byte'larını üret. Üretilemezse istisna fırlatır — `build_corpus` yakalar."""
    if case.scenario == "corrupt_ooxml":
        return _corrupt_ooxml(case.fmt, salt=len(case.case_id))
    if case.scenario == "zip_bomb":
        return _zip_bomb(case.fmt)
    if not case.planted_values:  # pragma: no cover — plan() her zaman en az bir değer koyar
        raise ValueError(f"{case.case_id}: yerleştirilecek değer yok")
    return _BUILDERS[case.fmt](case.scenario, case.planted_values[0], case.context)


def build_corpus() -> list[tuple[StressCase, bytes]]:
    """Tüm korpusu üret. Üretilemeyen vaka ATLANMAZ: `generated=False` + gerekçe ile listede
    kalır (byte'ları boş), böylece raporda görünür ve sessizce kaybolmaz."""
    out: list[tuple[StressCase, bytes]] = []
    for case in plan():
        try:
            data = build_one(case)
        except Exception as e:  # noqa: BLE001 — üretim hatası vakayı düşürmez, işaretler
            case.generated = False
            case.skip_reason = f"{type(e).__name__}: {e}"
            data = b""
        out.append((case, data))
    return out


def corpus_hashes() -> dict[str, str]:
    """case_id → sha256 (üretilemeyenler için boş string). Determinizm doğrulaması için."""
    return {
        c.case_id: (hashlib.sha256(d).hexdigest() if d else "")
        for c, d in build_corpus()
    }


def summary() -> dict:
    """Rapor-güvenli korpus özeti — ham değer içermez."""
    cases = [c for c, _ in build_corpus()]
    by_fmt: dict[str, int] = {}
    by_scenario: dict[str, int] = {}
    for c in cases:
        by_fmt[c.fmt] = by_fmt.get(c.fmt, 0) + 1
        by_scenario[c.scenario] = by_scenario.get(c.scenario, 0) + 1
    return {
        "total": len(cases),
        "by_format": by_fmt,
        "by_scenario": by_scenario,
        "generated": sum(1 for c in cases if c.generated),
        "not_generated": [
            {"case_id": c.case_id, "scenario": c.scenario, "reason": c.skip_reason}
            for c in cases if not c.generated
        ],
        "architecturally_safe": sum(1 for c in cases if c.architecturally_safe),
        "expected_fail_closed": sum(1 for c in cases if c.expected == FAIL_CLOSED),
        "expected_safe_output": sum(1 for c in cases if c.expected == SAFE_OUTPUT),
    }
