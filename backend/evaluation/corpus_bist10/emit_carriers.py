"""Render an `ExtractedContent` as PDF / DOCX / XLSX bytes carrying the SAME content.

Used only to build the balanced Part-B corpus (BENCHMARK_GUIDE.md §13): one real document's
content, re-emitted in three container formats, so a recall difference between formats can only be
attributed to format HANDLING, not to different content. Canary injection happens separately, on
these carrier bytes (`evaluation/corpus_bist10/inject.py` already handles pdf/docx/xlsx).

PDF reuses `app.export.render_pdf.render_content_pdf` directly — it already renders any
`ExtractedContent`, not just anonymized ones. The DOCX/XLSX emitters here are new and live in
`evaluation/`, not `app/export/`, on purpose: the product's real export path is deliberately
PDF-only for security reasons (see `render_pdf.py`'s docstring — "what was audited is exactly what
is shipped"); these are benchmark-corpus construction tooling, not a user-facing export feature.
"""
from __future__ import annotations

import io
import re
import zipfile

from app.export.render_pdf import render_content_pdf
from app.extraction.base import BlockType, ExtractedContent

# Sabit zaman damgası: benchmark taşıyıcıları BYTE-BYTE yeniden üretilebilir olmalı, yoksa
# "aynı manifest → aynı korpus" sözü tutmaz ve ikinci bir ekip (Codex) hash doğrulaması yapamaz.
# OOXML üreticileri (python-docx, openpyxl) hem ZIP girdi tarihlerine hem docProps/core.xml
# içine ÜRETİM ANINI yazar; aynı içerik iki farklı saniyede farklı sha256 verir. Bu ölçülmüştür:
# aynı içerik 2 sn arayla emit edildiğinde ZIP girdi tarihi (…,15,58,40) → (…,15,58,42).
_EPOCH = (1980, 1, 1, 0, 0, 0)
_FIXED_DT = "2026-01-01T00:00:00Z"
_CORE_DATE_RE = re.compile(
    rb"(<dcterms:(?:created|modified)[^>]*>)[^<]*(</dcterms:(?:created|modified)>)")


def _deterministic_ooxml(data: bytes) -> bytes:
    """OOXML (docx/xlsx) byte'larını yeniden üretilebilir hale getirir.

    İki kaynak birden normalize edilir — biri eksik kalırsa dosya yine değişir:
      1. ZIP girdi tarihleri → sabit epoch
      2. docProps/core.xml içindeki dcterms:created/modified → sabit tarih
    Girdi sırası ve sıkıştırma tipi korunur (aksi halde byte düzeni değişir).
    """
    src = zipfile.ZipFile(io.BytesIO(data))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as out:
        for item in src.infolist():
            payload = src.read(item.filename)
            if item.filename == "docProps/core.xml":
                payload = _CORE_DATE_RE.sub(
                    lambda m: m.group(1) + _FIXED_DT.encode() + m.group(2), payload)
            info = zipfile.ZipInfo(item.filename, date_time=_EPOCH)
            info.compress_type = item.compress_type
            info.external_attr = item.external_attr
            out.writestr(info, payload)
    return buf.getvalue()

# XML 1.0 forbids most C0 control chars (NUL, VT, FF, and the 0x0E-0x1F band); real BIST filings
# extracted via OCR/PDF-text occasionally carry these (observed: AKBNK annual reports), which makes
# python-docx raise ValueError and openpyxl raise IllegalCharacterError on write. Same char class
# openpyxl itself strips via `openpyxl.utils.exceptions.ILLEGAL_CHARACTERS_RE` — applied here too so
# python-docx (which has no built-in sanitizer) gets the same treatment.
_ILLEGAL_XML_CHARS_RE = re.compile(r"[\000-\010]|[\013-\014]|[\016-\037]")


def _clean(text: str) -> str:
    return _ILLEGAL_XML_CHARS_RE.sub("", text) if text else text


def render_content_docx(content: ExtractedContent) -> bytes:
    import docx as _docx

    d = _docx.Document()
    for b in content.blocks:
        if b.type == BlockType.heading and b.text:
            d.add_heading(_clean(b.text), level=2)
        elif b.type == BlockType.table and b.cells:
            rows = b.rows
            if not rows:
                continue
            t = d.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    t.rows[ri].cells[ci].text = _clean(val)
        elif b.text:
            d.add_paragraph(_clean(b.text))
    buf = io.BytesIO()
    d.save(buf)
    return _deterministic_ooxml(buf.getvalue())


def render_content_xlsx(content: ExtractedContent) -> bytes:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Content"
    row_idx = 1
    for b in content.blocks:
        if b.type == BlockType.heading and b.text:
            ws.cell(row=row_idx, column=1, value=_clean(b.text))
            row_idx += 2
        elif b.type == BlockType.table and b.cells:
            for row in b.rows:
                for ci, val in enumerate(row, start=1):
                    ws.cell(row=row_idx, column=ci, value=_clean(val))
                row_idx += 1
            row_idx += 1
        elif b.text:
            ws.cell(row=row_idx, column=1, value=_clean(b.text))
            row_idx += 1
    buf = io.BytesIO()
    wb.save(buf)
    return _deterministic_ooxml(buf.getvalue())


# PyMuPDF her kayıtta rastgele bir trailer /ID üretir. Ölçüldü: aynı içeriğin iki emit'i
# BOYUTÇA aynı, yalnız /ID'nin 62 byte'ı farklı — içerik birebir aynı. /ID belge kimliğidir,
# içerik değil; sabitlemek PDF'i geçersiz kılmaz. Trailer xref tablosundan SONRA geldiği için
# buradaki değişiklik hiçbir nesne offset'ini veya startxref'i kaydırmaz (üstelik uzunluk korunur).
_PDF_ID_RE = re.compile(rb"/ID\s*\[\s*<([0-9A-Fa-f]*)>\s*<([0-9A-Fa-f]*)>\s*\]")


def _deterministic_pdf(data: bytes) -> bytes:
    def _zero(m: bytes) -> bytes:
        return b"/ID[<" + b"0" * len(m.group(1)) + b"><" + b"0" * len(m.group(2)) + b">]"

    return _PDF_ID_RE.sub(_zero, data)


def _render_pdf_deterministic(content: ExtractedContent) -> bytes:
    return _deterministic_pdf(render_content_pdf(content))


_EMITTERS = {"pdf": _render_pdf_deterministic, "docx": render_content_docx,
             "xlsx": render_content_xlsx}


def emit(content: ExtractedContent, fmt: str) -> bytes:
    fn = _EMITTERS.get(fmt)
    if fn is None:
        raise ValueError(f"unsupported carrier format: {fmt!r} (expected pdf/docx/xlsx)")
    return fn(content)
