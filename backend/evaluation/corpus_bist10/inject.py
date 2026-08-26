"""Inject synthetic canaries into COPIES of real corpus documents, with exact ground truth.

Why into real documents: `evaluation/bist30/` planted canaries in synthetic carriers, so its recall
described a toy page. Here the same values are written into real Turkish/English investor-relations
files — real layout, real surrounding prose, real noise — so recall/precision describe the system's
behaviour on the documents it will actually meet.

Each canary is written next to a unique non-PII locator marker so it can be found afterwards, and
the values are the shared catalog from `evaluation/bist30/canary.py` (all synthetic).

PDF   → text lines in the bottom margin of up to N pages (the footer band: usually empty, and a
        realistic place for contact details in real reports).
DOCX  → body paragraph, table cell, header, footer, and a trailing paragraph block.
XLSX  → cells on the first sheet plus a fresh sheet.
"""
from __future__ import annotations

import io

from evaluation.bist30.canary import catalog_by_id
from evaluation.bist30.harness import Markers, Placement, make_placement

# Which canaries go into a document, and in which channel. Kept small per document so the injection
# does not dominate the page — the point is detection in context, not a wall of PII.
_PDF_PLAN = [("prs_tr", "text"), ("eml", "text"), ("gsm_tr", "text"), ("tckn", "text"),
             ("iban", "text"), ("acct", "text"), ("secret", "text"), ("adr", "text")]
_DOCX_PLAN = [("prs_tr", "body"), ("eml", "body"), ("iban", "table"), ("tckn", "header"),
              ("gsm_tr", "footer"), ("secret", "body"), ("acct", "body"), ("adr", "body")]
_XLSX_PLAN = [("prs_tr", "cell"), ("eml", "cell"), ("iban", "cell"), ("tckn", "cell"),
              ("acct", "cell"), ("secret", "hidden_sheet")]


def _plan(plan, mk: Markers, fmt: str) -> list[Placement]:
    cat = catalog_by_id()
    return [make_placement(mk, cat[cid], fmt, ch) for cid, ch in plan if cid in cat]


def inject_pdf(src: bytes, mk: Markers, max_pages: int = 3) -> tuple[bytes, list[Placement]]:
    import fitz

    from app.export.render_pdf import _FONTS_DIR
    font = str(_FONTS_DIR / "DejaVuSans.ttf")
    doc = fitz.open(stream=src, filetype="pdf")
    if doc.page_count == 0:
        doc.close()
        return src, []
    places = _plan(_PDF_PLAN, mk, "pdf")
    # Spread the canaries over the first, middle and last page so a partial extraction is visible.
    idx = sorted({0, doc.page_count // 2, doc.page_count - 1})[:max_pages]
    per_page = max(1, (len(places) + len(idx) - 1) // len(idx))
    for n, pno in enumerate(idx):
        page = doc[pno]
        chunk = places[n * per_page:(n + 1) * per_page]
        y = page.rect.height - 18 - 11 * len(chunk)   # bottom margin band
        for pl in chunk:
            try:
                page.insert_text((36, y), pl.carrier_text(), fontsize=8,
                                 fontfile=font, fontname="djv")
            except Exception:  # noqa: BLE001 — a hostile page must not abort the run
                pass
            y += 11
    out = doc.tobytes()
    doc.close()
    return out, places


def inject_docx(src: bytes, mk: Markers) -> tuple[bytes, list[Placement]]:
    import docx as _docx

    d = _docx.Document(io.BytesIO(src))
    places = _plan(_DOCX_PLAN, mk, "docx")
    by_ch: dict[str, list[Placement]] = {}
    for p in places:
        by_ch.setdefault(p.channel, []).append(p)

    for pl in by_ch.get("body", []):
        d.add_paragraph(pl.carrier_text())

    tbl_places = by_ch.get("table", [])
    if tbl_places:
        if d.tables:
            row = d.tables[0].add_row()
            for i, pl in enumerate(tbl_places):
                if i < len(row.cells):
                    row.cells[i].text = pl.carrier_text()
        else:  # no table in the source → make one so the channel is still exercised
            t = d.add_table(rows=1, cols=max(1, len(tbl_places)))
            for i, pl in enumerate(tbl_places):
                t.rows[0].cells[i].text = pl.carrier_text()

    for kind in ("header", "footer"):
        pls = by_ch.get(kind, [])
        if not pls or not d.sections:
            continue
        target = getattr(d.sections[0], kind)
        para = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
        para.text = " · ".join(p.carrier_text() for p in pls)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue(), places


def inject_xlsx(src: bytes, mk: Markers) -> tuple[bytes, list[Placement]]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(src))
    places = _plan(_XLSX_PLAN, mk, "xlsx")
    ws = wb.worksheets[0]
    row = ws.max_row + 3
    for i, pl in enumerate(p for p in places if p.channel == "cell"):
        ws.cell(row=row + i, column=1, value=pl.carrier_text())
    hidden = [p for p in places if p.channel == "hidden_sheet"]
    if hidden:
        hs = wb.create_sheet("Ek")
        for i, pl in enumerate(hidden, start=1):
            hs.cell(row=i, column=1, value=pl.carrier_text())
        hs.sheet_state = "hidden"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), places


_INJECTORS = {"pdf": inject_pdf, "docx": inject_docx, "xlsx": inject_xlsx}


def inject(src: bytes, fmt: str, mk: Markers) -> tuple[bytes, list[Placement]]:
    """Return (injected_bytes, ground_truth_placements); (src, []) if format unsupported."""
    fn = _INJECTORS.get(fmt)
    if fn is None:
        return src, []
    return fn(src, mk)
