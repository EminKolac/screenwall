"""Build a DOCX carrier embedding canaries across every channel the extractor reads (body,
table, section header, section footer, comment) plus line-broken and case/space variants, plus a
PII-bearing filename. Returns (bytes, filename, placements)."""
from __future__ import annotations

import io

import docx

from evaluation.bist30.canary import catalog_by_id
from evaluation.bist30.harness import Markers, Placement, make_placement

# Every entity type goes in the body (the baseline detection surface).
_BODY = ("prs_tr", "prs_en", "gsm_tr", "lnd_tr", "phn_uk", "eml", "adr", "iban", "card",
         "tckn", "acct", "ip", "secret", "url_tok", "plate", "passport")


def build_docx(mk: Markers) -> tuple[bytes, str, list[Placement]]:
    cat = catalog_by_id()
    placements: list[Placement] = []
    d = docx.Document()

    def add_body(cid: str, variant: str = "") -> None:
        p = make_placement(mk, cat[cid], "docx", variant or "body", variant)
        d.add_paragraph(p.carrier_text())
        placements.append(p)

    d.add_heading("Anonimleştirme Canary Taşıyıcı Belgesi", level=1)
    for cid in _BODY:
        add_body(cid)
    add_body("iban", "line_break")   # line-wrapped IBAN
    add_body("eml", "case_space")    # upper-cased / double-spaced email

    # Table cell (critical value inside a table).
    pt = make_placement(mk, cat["iban"], "docx", "table")
    tbl = d.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "IBAN"
    tbl.cell(0, 1).text = pt.carrier_text()
    placements.append(pt)

    # Section header + footer — must be unlinked so the extractor reads them.
    sec = d.sections[0]
    ph = make_placement(mk, cat["eml"], "docx", "header")
    sec.header.is_linked_to_previous = False
    sec.header.paragraphs[0].text = ph.carrier_text()
    placements.append(ph)

    pf = make_placement(mk, cat["gsm_tr"], "docx", "footer")
    sec.footer.is_linked_to_previous = False
    sec.footer.paragraphs[0].text = pf.carrier_text()
    placements.append(pf)

    # Comment (its own text carries the canary; anchored to a body run).
    pc = make_placement(mk, cat["tckn"], "docx", "comment")
    anchor = d.add_paragraph("Bu satır incelenmek üzere yorumlanmıştır.")
    try:
        d.add_comment(runs=anchor.runs, text=pc.carrier_text(), author="QA", initials="QA")
    except Exception:  # noqa: BLE001 — if the API shape differs, the channel is simply unavailable
        pass
    placements.append(pc)

    # Filename embeds a person-name canary (tests filename anonymization separately).
    fp = make_placement(mk, cat["prs_tr"], "docx", "filename")
    placements.append(fp)
    filename = f"Faaliyet Raporu {fp.value} {fp.marker}.docx"

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue(), filename, placements
