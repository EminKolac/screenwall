"""DOCX extraction via python-docx. Preserves body order (paragraphs + tables) and detects headings
by style. For full anonymization coverage it ALSO captures section headers/footers and comments so
PII living outside the body is audited + rendered too (Codex P1). Core-property metadata
(author/title) is captured but never re-emitted downstream.
"""
from __future__ import annotations

import io

import docx as pydocx
from docx.document import Document as _Doc
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.extraction.base import Block, BlockType, ExtractedContent, TableCell
from app.models.document import FileKind

_HEADER_ATTRS = ("header", "first_page_header", "even_page_header")
_FOOTER_ATTRS = ("footer", "first_page_footer", "even_page_footer")


def _iter_block_items(parent):
    parent_elm = parent.element.body if isinstance(parent, _Doc) else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _table_cells(table) -> list[TableCell]:
    return [
        TableCell(row=r, col=c, text=cell.text.strip())
        for r, row in enumerate(table.rows)
        for c, cell in enumerate(row.cells)
    ]


class DocxExtractor:
    kind = FileKind.docx

    def _extract_headers_footers(self, d, blocks: list[Block], start_id: int) -> int:
        """Header/footer paragraphs + tables, deduped (headers usually repeat across sections)."""
        bid = start_id
        seen: set[str] = set()
        for attrs, label in ((_HEADER_ATTRS, "header"), (_FOOTER_ATTRS, "footer")):
            for sec in d.sections:
                for attr in attrs:
                    part = getattr(sec, attr, None)
                    if part is None or getattr(part, "is_linked_to_previous", False):
                        continue
                    for p in part.paragraphs:
                        text = p.text.strip()
                        if text and text not in seen:
                            seen.add(text)
                            bid += 1
                            blocks.append(Block(block_id=f"{label}{bid}", type=BlockType.paragraph,
                                                text=text, location=label))
                    for tbl in part.tables:
                        cells = _table_cells(tbl)
                        key = f"{label}tbl:" + "|".join(c.text for c in cells)
                        if cells and key not in seen:
                            seen.add(key)
                            bid += 1
                            blocks.append(Block(block_id=f"{label}tbl{bid}", type=BlockType.table,
                                                cells=cells, location=f"{label} table"))
        return bid

    def extract(self, data: bytes, filename: str) -> ExtractedContent:
        d = pydocx.Document(io.BytesIO(data))
        blocks: list[Block] = []
        bid = 0

        # 1) Body — paragraphs (heading by style) + tables, in document order.
        for item in _iter_block_items(d):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                bid += 1
                style = (item.style.name if item.style else "") or ""
                is_head = style.startswith(("Heading", "Title"))
                btype = BlockType.heading if is_head else BlockType.paragraph
                blocks.append(Block(block_id=f"p{bid}", type=btype, text=text,
                                    location=f"para {bid}"))
            else:  # Table
                bid += 1
                blocks.append(Block(block_id=f"tbl{bid}", type=BlockType.table,
                                    cells=_table_cells(item), location=f"table {bid}"))

        # 2) Headers/footers (outside the body — never audited before v1).
        bid = self._extract_headers_footers(d, blocks, bid)

        # 3) Comments — each comment's text may itself carry PII.
        try:
            for cm in d.comments:
                text = (cm.text or "").strip()
                if text:
                    bid += 1
                    blocks.append(Block(block_id=f"cmt{bid}", type=BlockType.paragraph,
                                        text=text, location="comment"))
        except Exception:  # noqa: BLE001 — comments are best-effort; never fail extraction
            pass

        cp = d.core_properties
        meta = {k: v for k, v in {"author": cp.author or "", "title": cp.title or ""}.items() if v}
        return ExtractedContent(kind=FileKind.docx, blocks=blocks, metadata=meta)
